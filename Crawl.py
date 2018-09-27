import requests, sys, re, configparser
from docx import Document
from urllib.parse import urljoin
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

class Article(object):
    def __init__(self, title, link):
        self.title = title
        self.link = link
        self.category = ""
        self.author = ""
        self.date = ""
        self.content = ""


class SingleDay(object):
    def __init__(self, main_url):
        self.main_url = main_url
        self.urls = []
        self.articles = []

    def _getpage(self, url, css_rules):
        r = requests.get(url)
        r.raise_for_status()
        c = r.content
        soup = BeautifulSoup(c, "html.parser")
        results = soup.select(css_rules)
        return results

    # cssrules: .right_title-name a
    def get_urls(self, css_rules):
        pageList = self._getpage(self.main_url, css_rules)
        # from relative to absolute.
        self.urls = [urljoin(self.main_url, i.get('href')) for i in pageList]

    # cssrules1: .one a ; cssrules2:".text_c"
    def get_index(self, css_rules1, css_rules2):
        for url in self.urls:
            pageList = self._getpage(url, css_rules1)
            for i in pageList:
                s = str(i.contents)
                title = s.split('"')[1].strip()
                for s in keywords+subkeywords:
                    if s in title:
                        link = urljoin(self.main_url, i.get('href'))
                        article = Article(title, link)
                        if s in subkeywords:
                            article.category = "集锦"
                        else:
                            article.category = s

                        result = self._getpage(link, css_rules2)[0]
                        article.author = result.h4
                        article.date = ''.join(result.div.text.split())

                        paragraph = result.find(id='ozoom').find_all('p')
                        article.content = [c.text for c in paragraph]

                        self.articles.append(article)
                        break


    def write_to_docx(self):

        for i in self.articles:
            filename = i.category + '.docx'
            self.document = Document(filename)
            if "New Title" not in  self.document.styles:
                styles = self.document.styles
                new_heading_style = styles.add_style('New Title', WD_STYLE_TYPE.PARAGRAPH)
                new_heading_style.base_style = styles['Title']
                font = new_heading_style.font
                font.name = '宋体'
                font.size = Pt(18)

                #style: title:小二 宋体  paragrapth:三号 仿宋
            self.document.styles['Normal'].font.name = u'仿宋'

            self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            self.document.styles['New Title']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            self.document.styles['Normal'].font.size=Pt(16)



            head=self.document.add_paragraph(i.title,style="New Title")
            head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            author=self.document.add_paragraph(i.author)

            author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            date=self.document.add_paragraph(i.date)
            date.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for p in i.content:
                self.document.add_paragraph(p)
            self.document.save(filename)


if __name__ == "__main__":

    config = configparser.ConfigParser()
    config.read('config.ini', encoding='utf-8')
    keywords = config['DEFAULT']['Keywords'].split(',')
    subkeywords=config['DEFAULT']['SubKeywords'].split(',')
    period = {'year': config['DEFAULT']['Year'], 'month': config['DEFAULT']['Month'],
              'start': config['DEFAULT']['StartDay'], 'end': config['DEFAULT']['EndDay']}

    base_url = "http://paper.people.com.cn/rmrb/html/{year}-{month}/{day}/nbs.D110000renmrb_01.htm"


    for i in keywords+ ["集锦"]:
        doc = Document()
        doc.save(i + '.docx')

    for i in range(int(period['start']),int(period['end'])+1):
        main_url = base_url.format(year=period['year'], month=period['month'].zfill(2), day="%02d" % i)
        s = SingleDay(main_url)
        s.get_urls('.right_title-name a')
        s.get_index('#titleList a', ".text_c")
        s.write_to_docx()
